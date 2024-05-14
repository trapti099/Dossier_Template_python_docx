from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
import re
import warnings
warnings.filterwarnings("ignore")
def new_paragraph(obj,text,text_size=12,boolean_bold=False,boolean_italic=False,font_name ='Calibri',alignment = WD_ALIGN_PARAGRAPH.LEFT,r=0,g=0,b=0):
     text_para = obj.add_paragraph()
     run = text_para.add_run(text)
     run.font.size = Pt(text_size)
     run.font.color.rgb = RGBColor(r,g,b)
     run.font.bold = boolean_bold
     run.font.italic = boolean_italic
     run.font.name = font_name
     run.alignment = WD_ALIGN_PARAGRAPH.LEFT
     return text_para

def add_text_to_para(obj,text,text_size=12,boolean_bold=False,boolean_italic=False,font_name ='Calibri',alignment = WD_ALIGN_PARAGRAPH.LEFT,r=0,g=0,b=0):
     run = obj.add_run(text)
     run.font.size = Pt(text_size)
     run.font.color.rgb = RGBColor(r,g,b)
     run.font.bold = boolean_bold
     run.font.italic = boolean_italic
     run.font.name = font_name
     run.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_image_to_footer(document, image_path):
  section = document.sections[0]
  footer_img = section.footer  # Access the footer object
  section.left_margin = Cm(1)  #indent of the whole document from left becomes 1cm
  paragraph = footer_img.paragraphs[0]  
  run = paragraph.add_run()
  run.add_picture(image_path,width=Cm(19.35),height=Cm(1.09))  # Set image width

def format_table(table):
    # Format paragraphs in cells of the third column
    for cell in table.columns[2].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Format text in all cells of the table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    if run.text == ":":
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Set font color to red for ":"

# Function to add a row to the personal table
def add_row(table, left_text, middle_text, right_text):
    row = table.add_row().cells
    row[0].text = left_text
    row[1].text = middle_text
    row[2].text = right_text

# Function to add a row to the education table
def add_educational_row(table, year, institution, degree):
    row = table.add_row().cells
    row[0].text = year
    row[1].text = ":"
    institution_cell = row[2].paragraphs[0]
    institution_run = institution_cell.add_run(f"{institution}")
    institution_run.font.color.rgb = RGBColor(128, 0, 128)
    institution_cell.add_run(f"\n{degree}")

def add_additional_row(table, label, colon, company, info):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = colon
    cell = row[2]
    for paragraph in cell.paragraphs:
        cell._element.remove(paragraph._element)
    run_company = cell.add_paragraph().add_run(company)
    run_company.font.color.rgb = RGBColor(128, 0, 128)
    run_company.font.bold = True
    paragraph_info = cell.add_paragraph(info)

#def add_bullet_points(text,obj):
    #if text.find("*") != -1:
        #if text != None:
            #points = text.split('*')
            #points = [p.strip() for p in points if p.strip()]
            #for point in points:
		#obj.add_paragraph(point, style='ListBullet')
  	#else:
            #obj.add_paragraph()
  	

# def extract_responsibilities_and_achievements(text):

#     if text == None:
#         key_responsibilities='NA'
#         key_achievements = 'NA'
#     else:
#         responsibilities_pattern = re.compile(r'Main responsibilities(.*?)The (three )?most relevant results', re.DOTALL)
#         achievements_pattern = re.compile(r'The (three )?most relevant results(.*?)$', re.DOTALL)
#         responsibilities_match = responsibilities_pattern.search(text)
#         if responsibilities_match:
#             key_responsibilities = responsibilities_match.group(1).strip()
#         else:
#             key_responsibilities = None
#         achievements_match = achievements_pattern.search(text)
#         if achievements_match:
#             key_achievements = achievements_match.group(2).strip()
#         else:
#             key_achievements = None
#     return key_responsibilities, key_achievements
